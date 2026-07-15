"""
Exports the final report (goal + findings + report text) into a polished,
professional DOCX or PDF — not a raw JSON dump, an actually readable document
with a title page, headings, and a sources section.
"""

import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable
)
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

ACCENT_COLOR = "#4F46E5"  # indigo — used consistently across UI + exports

# ------------------------------------------------------------------
# Unicode font support for PDF export.
# ReportLab's default fonts (Helvetica etc.) are Latin-only — Hindi or
# Japanese text would silently render as blank boxes, not crash, which
# is worse (a silently broken deliverable). We register real Unicode
# fonts and pick one per-report based on the script actually used.
# ------------------------------------------------------------------
_FONTS_DIR = os.path.join(os.path.dirname(__file__), "fonts")
_DEVANAGARI_FONT = "NotoDevanagari"
_JAPANESE_FONT = "HeiseiKakuGo-W5"
_DEFAULT_FONT = "Helvetica"
_DEFAULT_FONT_BOLD = "Helvetica-Bold"

_fonts_registered = False


def _ensure_fonts_registered():
    global _fonts_registered
    if _fonts_registered:
        return
    try:
        registerFont(TTFont(_DEVANAGARI_FONT, os.path.join(_FONTS_DIR, "NotoSansDevanagari-Regular.ttf")))
    except Exception:  # noqa: BLE001
        pass  # falls back to default font; Devanagari text may not render, but export won't crash
    try:
        registerFont(UnicodeCIDFont(_JAPANESE_FONT))
    except Exception:  # noqa: BLE001
        pass
    _fonts_registered = True


def _pick_pdf_font(text: str) -> str:
    """Detects the dominant script in the text and returns the right font name."""
    if re.search(r"[\u0900-\u097F]", text):  # Devanagari block (Hindi)
        return _DEVANAGARI_FONT
    if re.search(r"[\u3040-\u30FF\u4E00-\u9FFF]", text):  # Hiragana/Katakana/Kanji (Japanese)
        return _JAPANESE_FONT
    return _DEFAULT_FONT


def _split_sections(report_text: str):
    """
    Splits the LLM's report text into (heading, body) pairs based on markdown
    '#' headers, so both DOCX and PDF exporters can render real structure
    instead of one giant paragraph.
    """
    lines = report_text.strip().split("\n")
    sections = []
    current_heading = None
    current_body = []

    for line in lines:
        stripped = line.strip()
        header_match = re.match(r"^#{1,3}\s+(.*)", stripped)
        if header_match:
            if current_heading is not None or current_body:
                sections.append((current_heading, "\n".join(current_body).strip()))
            current_heading = header_match.group(1)
            current_body = []
        else:
            current_body.append(line)

    if current_heading is not None or current_body:
        sections.append((current_heading, "\n".join(current_body).strip()))

    return sections


def export_docx(goal: str, report_text: str, sources: list, out_path: str):
    doc = Document()

    # Title
    title = doc.add_heading("TaskPilot Research Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(goal)
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated {datetime.now().strftime('%d %B %Y, %H:%M')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Body sections
    for heading, body in _split_sections(report_text):
        if heading:
            doc.add_heading(heading, level=1)
        if body:
            for para in body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

    # Sources
    if sources:
        doc.add_heading("Sources", level=1)
        for src in sources:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(src)

    doc.save(out_path)
    return out_path


def export_pdf(goal: str, report_text: str, sources: list, out_path: str):
    _ensure_fonts_registered()
    font_name = _pick_pdf_font(goal + " " + report_text)
    is_custom_font = font_name != _DEFAULT_FONT
    bold_font = font_name if is_custom_font else _DEFAULT_FONT_BOLD  # our TTF/CID fonts have no separate bold variant registered

    doc = SimpleDocTemplate(
        out_path, pagesize=letter,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "TitleCustom", parent=styles["Title"], fontSize=22, fontName=bold_font,
        textColor=HexColor(ACCENT_COLOR), spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleCustom", parent=styles["Normal"], fontSize=12, fontName=font_name,
        textColor=HexColor("#555555"), alignment=1, spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "MetaCustom", parent=styles["Normal"], fontSize=8, fontName=font_name,
        textColor=HexColor("#999999"), alignment=1, spaceAfter=20,
    )
    heading_style = ParagraphStyle(
        "HeadingCustom", parent=styles["Heading1"], fontSize=14, fontName=bold_font,
        textColor=HexColor("#1F2937"), spaceBefore=16, spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyCustom", parent=styles["Normal"], fontSize=10.5, fontName=font_name,
        leading=16, spaceAfter=10,
    )
    source_style = ParagraphStyle(
        "SourceCustom", parent=styles["Normal"], fontSize=9, fontName=font_name,
        textColor=HexColor("#4B5563"), leftIndent=12, spaceAfter=4,
    )

    story = []
    story.append(Paragraph("TaskPilot Research Report", title_style))
    story.append(Paragraph(goal, subtitle_style))
    story.append(Paragraph(f"Generated {datetime.now().strftime('%d %B %Y, %H:%M')}", meta_style))
    story.append(HRFlowable(width="100%", color=HexColor("#E5E7EB")))
    story.append(Spacer(1, 12))

    for heading, body in _split_sections(report_text):
        if heading:
            story.append(Paragraph(heading, heading_style))
        if body:
            for para in body.split("\n\n"):
                if para.strip():
                    safe = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe, body_style))

    if sources:
        story.append(Spacer(1, 10))
        story.append(Paragraph("Sources", heading_style))
        for src in sources:
            story.append(Paragraph(f"• {src}", source_style))

    doc.build(story)
    return out_path
